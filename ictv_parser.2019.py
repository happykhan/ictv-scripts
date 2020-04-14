#!/usr/bin/env python

# Copyright 2019 Nabil-Fareed Alikhan & Evelien Adriaenssens Licensed under the
#     Educational Community License, Version 2.0 (the "License"); you may
#     not use this file except in compliance with the License. You may
#     obtain a copy of the License at
#
#      http://www.osedu.org/licenses/ECL-2.0
#
#     Unless required by applicable law or agreed to in writing,
#     software distributed under the License is distributed on an "AS IS"
#     BASIS, WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express
#     or implied. See the License for the specific language governing
#     permissions and limitations under the License.

"""
Tabulates ICTV taxonomy proposals

Tabulates ICTV taxonomy proposals.

### CHANGE LOG ###
2019-06-25 Nabil-Fareed Alikhan <nabil@happykhan.com> & Evelien Adriaenssens
    * Initial build
"""

__licence__ = "ECL-2.0"
__author__ = "Nabil-Fareed Alikhan"
__author_email__ = "nabil@happykhan.com"
__version__ = "0.1.0"

epi = "Licence: " + __licence__ + " by " + __author__ + " <" + __author_email__ + ">"

from docx import Document
import logging
import os
import sys
import traceback
import time
import argparse

logger = logging.getLogger()
logger.setLevel(logging.INFO)


def get_meta(file_path):
    document = Document(file_path)
    values = {}
    values['application'] = os.path.basename(file_path)
    all_text = []
    for row in document.tables[0].rows:
        for cell in row.cells:
            if len(cell.tables) > 0 :
                for row in cell.tables[0].rows:
                    for cell in row.cells:
                        if cell.text not in all_text:
                            all_text.append(cell.text)
            if cell.text not in all_text:
                all_text.append(cell.text)
    count = 0
    values['code'] = all_text[all_text.index('Code assigned:') + 1].strip()
    values['title'] = [x for x in all_text if x.startswith('Short title:')][0].replace('\n', ' ').strip()
    values['title'] = values['title'].replace('Short title:', '')\
        .replace(' (e.g. “Create six new species in the genus Zetavirus”)' , '' )\
        .strip()
    try:
        values['authors'] = all_text[all_text.index('Provide email address for each author in a single line separated by semi-colons') + 1].replace('\n', ' ').strip()
        if values['authors'].endswith(','):
            values['authors'] = values['authors'][:-1]
        values['institution'] = all_text[all_text.index('Provide institutional addresses, each on a single line followed by author(s) initials (e.g. University of Woolloomooloo [SAB, HCL])') + 1].replace('\n', ';').strip()
        if values['institution'].endswith(';'):
            values['institution'] = values['institution'][:-1]
    except:
        logger.error('Could not parse authors or inst. for %s' % file_path)
    return values


def main():
    global args
    with open(args.output, 'w') as f:
        for app in os.listdir(args.app_dir):
            if app.endswith('.docx'):
                doc_path = os.path.join(args.app_dir, app)
                meta = get_meta(doc_path)
                f.write('\t'.join(meta.values()) + '\n')


if __name__ == "__main__":
    try:
        start_time = time.time()
        desc = __doc__.split("\n\n")[1].strip()
        parser = argparse.ArgumentParser(description=desc, epilog=epi)
        parser.add_argument(
            "-v", "--verbose", action="store_true", default=False, help="verbose output"
        )
        parser.add_argument(
            "--version", action="version", version="%(prog)s " + __version__
        )
        parser.add_argument(
            "-o", "--output", action="store", help="output prefix", default="summary.tsv"
        )
        parser.add_argument(
            "app_dir", action="store", help="Path to directory of applications docx"
        )
        args = parser.parse_args()
        if args.verbose:
            print("Executing @ " + time.asctime())
        main()
        if args.verbose:
            print("Ended @ " + time.asctime())
        sys.exit(0)
    except KeyboardInterrupt:  # Ctrl-C
        raise
    except SystemExit:  # sys.exit()
        raise
    except Exception:
        print("ERROR, UNEXPECTED EXCEPTION")
        print(traceback.print_exc())
        os._exit(1)
